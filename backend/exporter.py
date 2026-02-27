"""exporter.py — In-memory screenplay export helpers for PDF, DOCX, and TXT.

All functions return raw bytes and never write to disk.
"""

from __future__ import annotations

import io
import logging
from typing import List, Tuple

logger: logging.Logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Screenplay parsing helpers
# ---------------------------------------------------------------------------

_SCENE_PREFIXES: Tuple[str, ...] = ("INT.", "EXT.")


def _is_scene_heading(line: str) -> bool:
    stripped: str = line.strip()
    return any(stripped.startswith(prefix) for prefix in _SCENE_PREFIXES)


def _is_character_name(
    stripped: str,
    prev_was_blank: bool,
    next_nonblank: str,
) -> bool:
    """Return True if the line looks like a screenplay character name.

    Criteria:
        - Non-empty
        - Entirely upper-case
        - Previous line was blank (or it's the start of the script)
        - The next non-blank line is NOT a scene heading
    """
    if not stripped:
        return False
    if stripped != stripped.upper():
        return False
    if not prev_was_blank:
        return False
    if _is_scene_heading(next_nonblank):
        return False
    # Scene headings themselves would have already been caught above;
    # exclude them explicitly to avoid double-marking.
    if _is_scene_heading(stripped):
        return False
    return True


def _classify_lines(screenplay: str) -> List[Tuple[str, str]]:
    """Parse screenplay text into a list of (type, text) tuples.

    Types: 'scene', 'character', 'dialogue', 'action', 'blank'
    """
    raw_lines: List[str] = screenplay.split("\n")
    results: List[Tuple[str, str]] = []

    # Build a look-ahead list of "next non-blank line" for each position.
    next_nonblank: List[str] = [""] * len(raw_lines)
    last_nonblank: str = ""
    for i in range(len(raw_lines) - 1, -1, -1):
        if raw_lines[i].strip():
            last_nonblank = raw_lines[i].strip()
        next_nonblank[i] = last_nonblank

    in_dialogue: bool = False
    prev_was_blank: bool = True  # treat start-of-script as "after blank"

    for i, raw_line in enumerate(raw_lines):
        stripped: str = raw_line.strip()

        if not stripped:
            results.append(("blank", ""))
            in_dialogue = False
            prev_was_blank = True
            continue

        if _is_scene_heading(stripped):
            results.append(("scene", stripped))
            in_dialogue = False
            prev_was_blank = False
            continue

        if in_dialogue:
            results.append(("dialogue", stripped))
            prev_was_blank = False
            continue

        nxt: str = next_nonblank[i + 1] if i + 1 < len(raw_lines) else ""
        if _is_character_name(stripped, prev_was_blank, nxt):
            results.append(("character", stripped))
            in_dialogue = True
            prev_was_blank = False
            continue

        results.append(("action", stripped))
        prev_was_blank = False

    return results


# ---------------------------------------------------------------------------
# PDF export
# ---------------------------------------------------------------------------

def export_pdf(screenplay: str, project_id: str) -> bytes:
    """Generate a PDF from *screenplay* and return its bytes.

    Args:
        screenplay: Raw screenplay text.
        project_id: Used only for logging / metadata.

    Returns:
        bytes representing the PDF file.

    Raises:
        Exception: Re-raises any ReportLab / IO error so the caller can
            convert it into an HTTP error response.
    """
    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.units import inch
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.platypus import (
            SimpleDocTemplate,
            Paragraph,
            Spacer,
        )
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
        from reportlab.platypus import PageTemplate, Frame
        from reportlab.platypus.frames import Frame as LayoutFrame
        from reportlab.lib import colors

        buf: io.BytesIO = io.BytesIO()

        # Page layout
        page_width, page_height = LETTER
        margin_top: float = 1.0 * inch
        margin_bottom: float = 1.0 * inch
        margin_left: float = 1.5 * inch
        margin_right: float = 1.0 * inch

        # --- Page-number callback ----------------------------------------
        def _add_page_number(canvas, doc) -> None:  # type: ignore[override]
            page_num: int = canvas.getPageNumber()
            if page_num > 1:
                canvas.saveState()
                canvas.setFont("Courier", 12)
                canvas.drawRightString(
                    page_width - margin_right,
                    page_height - 0.5 * inch,
                    str(page_num),
                )
                canvas.restoreState()

        doc = SimpleDocTemplate(
            buf,
            pagesize=LETTER,
            topMargin=margin_top,
            bottomMargin=margin_bottom,
            leftMargin=margin_left,
            rightMargin=margin_right,
        )

        # Paragraph styles — Courier 12 throughout
        _font_name: str = "Courier"
        _font_size: int = 12
        _leading: int = 14

        style_scene = ParagraphStyle(
            "scene",
            fontName="Courier-Bold",
            fontSize=_font_size,
            leading=_leading,
            alignment=TA_LEFT,
            spaceAfter=6,
            spaceBefore=12,
        )
        style_character = ParagraphStyle(
            "character",
            fontName=_font_name,
            fontSize=_font_size,
            leading=_leading,
            alignment=TA_CENTER,
            spaceBefore=12,
        )
        style_dialogue = ParagraphStyle(
            "dialogue",
            fontName=_font_name,
            fontSize=_font_size,
            leading=_leading,
            alignment=TA_LEFT,
            leftIndent=1.5 * inch,
            rightIndent=1.0 * inch,
        )
        style_action = ParagraphStyle(
            "action",
            fontName=_font_name,
            fontSize=_font_size,
            leading=_leading,
            alignment=TA_LEFT,
            spaceAfter=6,
        )

        classified: List[Tuple[str, str]] = _classify_lines(screenplay)
        story = []

        for kind, text in classified:
            if kind == "blank":
                story.append(Spacer(1, 6))
                continue
            # Escape all XML special characters ReportLab's Paragraph parser requires
            safe_text: str = (
                text.replace("&", "&amp;")
                    .replace("<", "&lt;")
                    .replace(">", "&gt;")
                    .replace('"', "&quot;")
                    .replace("'", "&apos;")
            )
            if kind == "scene":
                story.append(Paragraph(safe_text.upper(), style_scene))
            elif kind == "character":
                story.append(Paragraph(safe_text, style_character))
            elif kind == "dialogue":
                story.append(Paragraph(safe_text, style_dialogue))
            else:  # action
                story.append(Paragraph(safe_text, style_action))

        doc.build(story, onFirstPage=_add_page_number, onLaterPages=_add_page_number)
        logger.info("export_pdf: generated PDF for project_id=%s (%d bytes)", project_id, buf.tell())
        return buf.getvalue()

    except Exception as exc:
        logger.error("export_pdf failed for project_id=%s: %s", project_id, exc)
        raise


# ---------------------------------------------------------------------------
# DOCX export
# ---------------------------------------------------------------------------

def export_docx(screenplay: str, project_id: str) -> bytes:
    """Generate a DOCX from *screenplay* and return its bytes.

    Args:
        screenplay: Raw screenplay text.
        project_id: Used only for logging / metadata.

    Returns:
        bytes representing the DOCX file.

    Raises:
        Exception: Re-raises any python-docx / IO error.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Set page margins on the single default section
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1)

        # Remove the blank default paragraph if present
        if doc.paragraphs:
            # Access underlying XML to delete the default empty paragraph
            p_elem = doc.paragraphs[0]._element
            p_elem.getparent().remove(p_elem)

        classified: List[Tuple[str, str]] = _classify_lines(screenplay)

        def _set_font(run, bold: bool = False) -> None:
            run.font.name = "Courier New"
            run.font.size = Pt(12)
            run.bold = bold

        for kind, text in classified:
            if kind == "blank":
                # Add an empty paragraph to represent blank lines
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                continue

            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)

            if kind == "scene":
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(text.upper())
                _set_font(run, bold=True)

            elif kind == "character":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                _set_font(run)

            elif kind == "dialogue":
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.left_indent = Inches(1.5)
                run = p.add_run(text)
                _set_font(run)

            else:  # action
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(text)
                _set_font(run)

        buf: io.BytesIO = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        logger.info("export_docx: generated DOCX for project_id=%s (%d bytes)", project_id, len(buf.getvalue()))
        return buf.getvalue()

    except Exception as exc:
        logger.error("export_docx failed for project_id=%s: %s", project_id, exc)
        raise


# ---------------------------------------------------------------------------
# TXT export
# ---------------------------------------------------------------------------

def export_txt(screenplay: str) -> bytes:
    """Return *screenplay* encoded as UTF-8 bytes.

    Args:
        screenplay: Raw screenplay text.

    Returns:
        UTF-8 encoded bytes of the screenplay.

    Raises:
        Exception: Re-raises any encoding error.
    """
    try:
        result: bytes = screenplay.encode("utf-8")
        logger.info("export_txt: encoded %d characters to UTF-8", len(screenplay))
        return result
    except Exception as exc:
        logger.error("export_txt failed: %s", exc)
        raise
